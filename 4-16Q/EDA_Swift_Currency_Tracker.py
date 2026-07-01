import pandas as pd
import missingno as msno
import matplotlib.pyplot as plt
import seaborn as sns
import os

from docx import Document
from docx.shared import Inches

# =========================================
# CREATE FOLDERS
# =========================================

os.makedirs("charts", exist_ok=True)

# =========================================
# LOAD DATASET
# =========================================

df = pd.read_csv("swift_currency_tracker_all_reports-1.csv")

# =========================================
# BASIC INFORMATION
# =========================================

print("\n===================================")
print("FIRST 5 ROWS")
print("===================================")

print(df.head())

print("\n===================================")
print("COLUMN NAMES")
print("===================================")

print(df.columns)

print("\n===================================")
print("DATASET INFO")
print("===================================")

df.info()

print("\n===================================")
print("SUMMARY STATISTICS")
print("===================================")

print(df.describe())

# =========================================
# DATA QUALITY CHECKS
# =========================================

print("\n===================================")
print("MISSING VALUES")
print("===================================")

print(df.isnull().sum())

print("\n===================================")
print("DUPLICATE ROWS")
print("===================================")

print(df.duplicated().sum())

# =========================================
# MISSING VALUE VISUALIZATION
# =========================================

msno.matrix(df)

plt.title("Missing Values Matrix")

plt.savefig("charts/missing_values_matrix.png")

plt.close()

# =========================================
# DISTRIBUTION ANALYSIS
# =========================================

plt.figure(figsize=(10, 5))

sns.histplot(df['value'], kde=True)

plt.title("Distribution of Currency Values")
plt.xlabel("Value")
plt.ylabel("Frequency")

plt.savefig("charts/histogram.png")

plt.close()

# =========================================
# BOXPLOT ANALYSIS
# =========================================

plt.figure(figsize=(10, 5))

sns.boxplot(x=df['value'])

plt.title("Boxplot of Currency Values")

plt.savefig("charts/boxplot.png")

plt.close()

# =========================================
# TOP 10 ECONOMIES BY VALUE
# =========================================

top10 = df.sort_values(
    by='value',
    ascending=False
).head(10)

plt.figure(figsize=(12, 6))

sns.barplot(
    data=top10,
    x='currency_or_economy',
    y='value'
)

plt.xticks(rotation=45)

plt.title("Top 10 Economies by Value")

plt.xlabel("Currency / Economy")
plt.ylabel("Value")

plt.tight_layout()

plt.savefig("charts/top10_economies.png")

plt.close()

# =========================================
# RMB GLOBAL RANK DISTRIBUTION
# =========================================

if 'rmb_global_rank' in df.columns:

    plt.figure(figsize=(10, 5))

    sns.histplot(
        df['rmb_global_rank'].dropna(),
        kde=True
    )

    plt.title("Distribution of RMB Global Rank")

    plt.xlabel("RMB Global Rank")
    plt.ylabel("Frequency")

    plt.savefig("charts/rmb_rank_distribution.png")

    plt.close()

# =========================================
# CORRELATION HEATMAP
# =========================================

corr = df.corr(numeric_only=True)

plt.figure(figsize=(8, 5))

sns.heatmap(
    corr,
    annot=True,
    cmap="coolwarm"
)

plt.title("Correlation Heatmap")

plt.savefig("charts/heatmap.png")

plt.close()

# =========================================
# OUTLIER DETECTION
# =========================================

Q1 = df['value'].quantile(0.25)
Q3 = df['value'].quantile(0.75)

IQR = Q3 - Q1

lower_bound = Q1 - 1.5 * IQR
upper_bound = Q3 + 1.5 * IQR

outliers = df[
    (df['value'] < lower_bound) |
    (df['value'] > upper_bound)
]

print("\n===================================")
print("OUTLIERS DETECTED")
print("===================================")

print(outliers)

print("\nNumber of Outliers:", len(outliers))

# =========================================
# SAVE CLEANED DATA
# =========================================

df.to_csv(
    "cleaned_currency_data.csv",
    index=False
)

print("\n===================================")
print("EDA COMPLETED SUCCESSFULLY")
print("===================================")

print("\nCharts saved inside 'charts' folder.")
print("Cleaned dataset saved as 'cleaned_currency_data.csv'")

# =========================================
# CREATE DOCX REPORT
# =========================================

doc = Document()

# =========================================
# TITLE
# =========================================

doc.add_heading(
    'Swift Currency Tracker - EDA Report',
    level=0
)

# =========================================
# OVERVIEW
# =========================================

doc.add_heading('1. Overview', level=1)

doc.add_paragraph(
    f"""
The dataset contains {len(df)} rows and {len(df.columns)} columns.

This exploratory data analysis (EDA) report examines
currency/economy values, RMB rankings, distributions,
correlations, and outlier behaviour.
"""
)

# =========================================
# DATA QUALITY
# =========================================

doc.add_heading('2. Data Quality Checks', level=1)

missing_values = df.isnull().sum().sum()
duplicates = df.duplicated().sum()

doc.add_paragraph(
    f"""
Missing Values: {missing_values}

Duplicate Rows: {duplicates}

The dataset was checked for completeness,
missing values, and duplicate records.
"""
)

doc.add_picture(
    "charts/missing_values_matrix.png",
    width=Inches(6)
)

# =========================================
# DISTRIBUTION ANALYSIS
# =========================================

doc.add_heading('3. Distribution Analysis', level=1)

doc.add_paragraph(
    """
Histogram and boxplot analysis were performed
to understand the spread, skewness,
and distribution of currency values.
"""
)

doc.add_picture(
    "charts/histogram.png",
    width=Inches(6)
)

doc.add_picture(
    "charts/boxplot.png",
    width=Inches(6)
)

# =========================================
# TOP ECONOMIES
# =========================================

doc.add_heading('4. Top Economies by Value', level=1)

doc.add_paragraph(
    """
The top 10 economies/currencies were identified
based on the highest recorded values.
"""
)

doc.add_picture(
    "charts/top10_economies.png",
    width=Inches(6)
)

# =========================================
# RMB RANK ANALYSIS
# =========================================

if 'rmb_global_rank' in df.columns:

    doc.add_heading(
        '5. RMB Global Rank Analysis',
        level=1
    )

    doc.add_paragraph(
        """
Distribution analysis of RMB global rankings
was performed to understand ranking patterns
across economies.
"""
    )

    doc.add_picture(
        "charts/rmb_rank_distribution.png",
        width=Inches(6)
    )

# =========================================
# CORRELATION ANALYSIS
# =========================================

doc.add_heading(
    '6. Correlation Analysis',
    level=1
)

doc.add_paragraph(
    """
Correlation heatmap analysis was used
to identify relationships between
numerical variables in the dataset.
"""
)

doc.add_picture(
    "charts/heatmap.png",
    width=Inches(6)
)

# =========================================
# OUTLIER ANALYSIS
# =========================================

doc.add_heading(
    '7. Outlier Detection',
    level=1
)

doc.add_paragraph(
    f"""
Outlier detection using the IQR method
identified {len(outliers)} potential
outliers in the dataset.
"""
)

# =========================================
# KEY TAKEAWAYS
# =========================================

doc.add_heading(
    '8. Key Takeaways',
    level=1
)

doc.add_paragraph(
    """
• The dataset was successfully cleaned and analyzed.

• Distribution analysis revealed the spread
  and skewness of currency values.

• Correlation analysis highlighted relationships
  between numerical variables.

• Outlier detection identified unusual
  observations requiring further investigation.

• Visualizations provide insights into
  economy rankings and RMB positioning.
"""
)
# =========================================
# OPEN RESEARCH QUESTIONS
# =========================================

doc.add_heading(
    '9. Open Research Questions',
    level=1
)

doc.add_paragraph(
    """
Q1. Are there significant relationships between RMB global rankings
and currency/economy values?

Q2. Which economies consistently exhibit unusually high or low
currency values, and what factors may explain these outliers?

Q3. How can additional macroeconomic indicators improve
future currency and economy analysis?
"""
)
# =========================================
# SAVE REPORT
# =========================================

doc.save(
    "Swift_Currency_Tracker_EDA_Report.docx"
)

print("\n===================================")
print("DOCX REPORT GENERATED SUCCESSFULLY")
print("===================================")

print("\nSaved as:")
print("Swift_Currency_Tracker_EDA_Report.docx")